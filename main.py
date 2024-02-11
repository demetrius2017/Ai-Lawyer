import json
import logging
import os
import re

import openai
from docx import Document
from dotenv import load_dotenv

# Загружаем переменные окружения из файла .env
load_dotenv()

logger = logging.getLogger(__name__)

# Теперь вы можете безопасно получить ваш API ключ
openai_api_key = os.getenv("Api")
if openai_api_key is None:
    raise ValueError("API ключ не найден. Убедитесь, что вы правильно установили переменную 'Api' в файле .env.")
else:
    openai.api_key = openai_api_key

# Глобальная переменная для хранения логгера
logger = None


def set_logger(custom_logger):
    global logger
    logger = custom_logger


def read_document(filename):
    """Чтение документа и возврат текста."""
    doc = Document(filename)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])


def analyze_document_from_perspective(document_text, perspective):
    """
    ... (rest of the function) ...
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-0125",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": f"Вы высококвалифицированный ИИ, действующий как {perspective}. Прочитайте следующий юридический документ и выделите ключевые пункты, по которым могут возникнуть разногласия. Укажи четкий номер пункта и причину",
            },
            {
                "role": "user",
                "content": document_text,
            },
        ],
    )
    return response.choices[0].message.content


def load_and_split_contract(filename):
    doc = Document(filename)
    # Регулярное выражение для поиска шаблонов типа "1.1", "2.1", "a)", "1)", "1.1.11" и т.д.
    pattern = re.compile(r"\b(\d+\.\d+|\d+\.\d+\.\d+|[a-z]\)|\d+\))")

    clauses = []
    for para in doc.paragraphs:
        if para.text.strip() != "" and pattern.match(para.text.strip()):
            clauses.append(para.text)
    return clauses


def get_feedback_from_gpt4(clause, perspective, feedback):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-0125",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Вы высококвалифицированный ИИ, действующий как юрист со специализацией в понимании языка, юридическом анализе и суммировании документов. Ваша задача - прочитать следующий юридический документ и предоставить краткое содержание его основных пунктов, включая любые потенциальные юридические вопросы или соображения. Сосредоточьтесь на наиболее критических аспектах, чтобы предложить с coherent и читаемое резюме, которое позволит кому-то понять ключевые юридические моменты без необходимости читать весь документ.",
            },
            {
                "role": "user",
                "content": f"С точки зрения {perspective}, как можно улучшить этот пунк Пункт: '{clause}' учитывая замечания {feedback}",
            },
        ],
    )
    return response.choices[0].message.content


def synthesize_compromise(clause, part1_feedback, part2_feedback):
    """
    Создает компромиссное предложение на основе обратной связи от двух сторон.

    Args:
        clause: Оригинальный текст пункта договора.
        part1_feedback: Обратная связь от первой стороны.
        part2_feedback: Обратная связь от второй стороны.

    Returns:
        Текст компромиссного предложения.
    """
    synthesis_prompt = f"Создайте сбалансированное предложение для пункта договора, учитывая следующие предложения от сторон: \nПункт договора: '{clause}'\nПредложение от первой стороны: {part1_feedback}\nПредложение от второй стороны: {part2_feedback}\nКомпромиссное предложение должно учитывать интересы обеих сторон и быть юридически корректным."

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo-0125",
            messages=[
                {
                    "role": "system",
                    "content": "Вы юрист, специализирующийся на согласовании договоров. Ваша задача - сформулировать компромиссное предложение.",
                },
                {"role": "user", "content": synthesis_prompt},
            ],
        )

        # Извлечение и возврат компромиссного предложения
        compromise = response.choices[0].message["content"].strip()
        return compromise

    except Exception as e:
        print(f"Произошла ошибка при создании компромиссного предложения: {e}")
        return "Не удалось сформулировать компромиссное предложение."


def negotiate_clauses(parties_info, clauses, executor_analysis, customer_analysis):
    """
    Функция применения сторон и выработки компромисса.

    Args:
        clauses: Список пунктов договора.
        executor_analysis: Анализ с точки зрения "Исполнителя".
        customer_analysis: Анализ с точки зрения "Заказчика".

    Returns:
        Список пунктов договора с итоговым вариантом, по которому обе стороны согласны.
    """
    logger.telega(f"Составляем компромисное соглашение...")

    part1 = parties_info["part1"]
    part2 = parties_info["part2"]

    final_clauses = []

    for clause in clauses:
        logger.telega(f"Пункт: {clause}")

        # Получаем предложения по улучшению пункта от "Исполнителя"
        part1_feedback = get_feedback_from_gpt4(clause, part1, f"основываясь на анализе {part1}")
        logger.telega(f"{part1} предложение: {part1_feedback}")

        # Получаем предложения по улучшению пункта от "Заказчика"
        part2_feedback = get_feedback_from_gpt4(clause, part2, f"основываясь на анализе {part2}")
        logger.telega(f"Заказчик предложение: {part2_feedback}")

        # Для примера, допустим, что финальная версия пункта формируется как совет от исполнителя
        final_clause = synthesize_compromise(clause, part1_feedback, part2_feedback)

        final_clauses.append({"original": clause, "final": final_clause})

        return final_clauses

    return final_clauses


def identify_parties(document_text, attempts=3):
    """
    Анализирует текст документа и идентифицирует стороны договора, возвращая информацию в JSON.
    Повторяет попытку в случае неудачи до заданного количества раз.

    :param document_text: Текст документа для анализа.
    :param attempts: Количество попыток получить корректный ответ.
    :return: Словарь с идентифицированными сторонами.
    """
    logger.telega("Идентификация сторон договора...")
    for attempt in range(attempts):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4-0125-preview",
                messages=[
                    {
                        "role": "system",
                        "content": "Определите стороны в договоре, указанные в тексте ниже, и верните результат в формате JSON, включая возможные роли, такие как покупатель и продавец, арендатор и арендодатель, заказчик и исполнитель. СТРОГИЙ формата ответа: {'part1': 'Заказчик', 'part2': 'Исполнитель'}. Обязательно соблюдение четких переменных part1 и part2.",
                    },
                    {"role": "user", "content": document_text[:500]},
                ],
            )

            # Преобразуем ответ от GPT в JSON
            response_text = response.choices[0].message["content"].strip()

            # Попытка преобразовать текстовый ответ в словарь Python
            parties_dict = json.loads(response_text)

            # Проверка наличия ключей part1 и part2 в ответе
            if "part1" in parties_dict and "part2" in parties_dict:
                logger.telega(f"Идентификация сторон договора завершена: {parties_dict}")
                return parties_dict
            else:
                logger.telega(f"Попытка {attempt + 1}: Ответ не содержит ожидаемых ключей 'part1' и 'part2'.")
        except Exception as e:
            logger.telega(f"Попытка {attempt + 1}: Произошла ошибка при идентификации сторон договора: {e}")

    logger.telega("Не удалось идентифицировать стороны договора после нескольких попыток.")
    return {}


def save_document(parties_info, filename, clauses):
    # Создание каталога, если он не существует
    os.makedirs(os.path.dirname(filename), exist_ok=True)

    part1 = parties_info["part1"]
    part2 = parties_info["part2"]

    doc = Document()
    for clause in clauses:
        doc.add_paragraph(clause["original"])

        # Добавляем обратную связь от первой стороны
        doc.add_paragraph(f"{part1} Feedback: {clause.get('part1_feedback', 'Нет фидбека')}")

        # Добавляем обратную связь от второй стороны
        doc.add_paragraph(f"{part2} Feedback: {clause.get('part2_feedback', 'Нет фидбека')}")

        # Добавляем финальную версию пункта
        doc.add_paragraph(f"Финальная версия: {clause.get('final', 'Финальная версия не определена')}")

    # Сохраняем документ
    doc.save(filename)


def process_document(input_file_path="downloads/contract_software_development.docx"):
    report_filename = "./sent/processed_" + os.path.basename(input_file_path)

    document_text = read_document(input_file_path)
    parties_info = identify_parties(document_text)
    print(parties_info)
    part1 = parties_info["part1"]
    part2 = parties_info["part2"]

    # Анализируем документ с точки зрения Исполнителя и Заказчика
    executor_analysis = analyze_document_from_perspective(document_text, part1)
    logger.telega(f"Анализ документа с точки зрения {part1}: {executor_analysis}")
    customer_analysis = analyze_document_from_perspective(document_text, part2)
    logger.telega(f"Анализ документа с точки зрения {part2}: {customer_analysis}")
    # Загружаем и разделяем документ на пункты
    clauses = load_and_split_contract(input_file_path)

    # Проводим переговоры по пунктам
    negotiated_clauses = negotiate_clauses(parties_info, clauses, executor_analysis, customer_analysis)

    # Сохраняем документ с итоговыми пунктами
    save_document(parties_info, report_filename, negotiated_clauses)

    return report_filename


def main():
    filename = "downloads/contract_software_development.docx"
    process_document(filename)


if __name__ == "__main__":
    main()
